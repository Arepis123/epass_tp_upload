"""Generate ePass TP Upload User Tutorial in the same style as PayslipSplitter_Tutorial.docx"""

from docx import Document
from docx.shared import Pt, RGBColor, Twips, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ---------- colour constants ----------
NAVY   = RGBColor(0x0A, 0x3D, 0x62)
LBLUE  = RGBColor(0xA8, 0xC5, 0xDA)
RED    = RGBColor(0xC0, 0x39, 0x2B)
GRAY   = RGBColor(0x55, 0x55, 0x55)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF0, 0xF2, 0xF5)
GREEN  = RGBColor(0xE8, 0xF5, 0xE9)

NAVY_HEX  = "0A3D62"
LGRAY_HEX = "F0F2F5"
WHITE_HEX = "FFFFFF"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_bottom_border(paragraph, color_hex, sz=6):
    """Add a bottom border line under a paragraph (used for PART headings)."""
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    str(sz * 8))   # sz in eighths of a point
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color_hex)
    pBdr.append(bot)
    pPr.append(pBdr)


def set_spacing(paragraph, before=0, after=0):
    pPr  = paragraph._p.get_or_add_pPr()
    pSpc = pPr.find(qn("w:spacing"))
    if pSpc is None:
        pSpc = OxmlElement("w:spacing")
        pPr.append(pSpc)
    if before:
        pSpc.set(qn("w:before"), str(before))
    if after:
        pSpc.set(qn("w:after"),  str(after))


def add_part_heading(doc, text):
    """Big section heading with bottom border (e.g. PART 1 — Title)."""
    p = doc.add_paragraph()
    set_spacing(p, before=280, after=40)
    add_bottom_border(p, NAVY_HEX, sz=6)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    run.font.size = Pt(13)
    return p


def add_step_heading(doc, text):
    """Step N: Description heading."""
    p = doc.add_paragraph()
    set_spacing(p, before=160, after=40)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = NAVY
    return p


def add_bullet(doc, text, bold_parts=None):
    """
    Add a bullet paragraph.
    bold_parts: list of substrings to render bold within the bullet text.
    """
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=40)
    if bold_parts:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                p.add_run(remaining[:idx])
            br = p.add_run(bp)
            br.bold = True
            remaining = remaining[idx + len(bp):]
        if remaining:
            p.add_run(remaining)
    else:
        p.add_run(text)
    return p


def add_code_block(doc, lines):
    """Monospaced code/config block."""
    p = doc.add_paragraph()
    set_spacing(p, after=40)
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "454")
    pPr.append(ind)
    first = True
    for line in lines:
        if not first:
            br = OxmlElement("w:br")
            p.runs[-1]._r.addnext(br) if p.runs else p._p.append(br)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        first = False
    return p


def add_warning(doc, text):
    """Red italic warning paragraph."""
    p = doc.add_paragraph()
    set_spacing(p, after=40)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RED
    run.font.size = Pt(10)
    return p


def add_page_break(doc):
    p  = doc.add_paragraph()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    p._p.append(br)


# ── document ─────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Twips(1134)
section.bottom_margin = Twips(1134)
section.left_margin   = Twips(1417)
section.right_margin  = Twips(1417)

# ── COVER TABLE ──────────────────────────────────────────────────────────────
cover = doc.add_table(rows=1, cols=1)
cover.style = "Table Grid"
cell = cover.cell(0, 0)
set_cell_bg(cell, NAVY_HEX)

# Row height
tr   = cover.rows[0]._tr
trPr = tr.get_or_add_trPr()
trH  = OxmlElement("w:trHeight")
trH.set(qn("w:val"),  "1081")
trH.set(qn("w:hRule"), "atLeast")
trPr.append(trH)

# Cell width — full page (auto)
tcPr = cell._tc.get_or_add_tcPr()
tcW  = OxmlElement("w:tcW")
tcW.set(qn("w:w"),    "9588")
tcW.set(qn("w:type"), "dxa")
tcPr.append(tcW)

def set_cover_para(p_obj, text, sz, bold=False, italic=False, color=WHITE):
    p_obj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_obj.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.size      = Pt(sz)

# First paragraph already exists in the cell
p0 = cell.paragraphs[0]
set_cover_para(p0, "ePass TP Upload", 40, bold=True, color=WHITE)

p1 = cell.add_paragraph()
set_cover_para(p1, "Construction Labour Exchange Centre Berhad", 10, color=LBLUE)

p2 = cell.add_paragraph()
set_cover_para(p2, "User Tutorial", 12, bold=True, italic=True, color=WHITE)

# Spacer after cover
doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# PART 1
# ═══════════════════════════════════════════════════════════════════════════
add_part_heading(doc, "PART 1 — First Time Setup (IT Department Setup)")

add_step_heading(doc, "Step 1: Install Prerequisites")
add_bullet(doc, "Ensure Python 3.9 or later is installed on the workstation.")
add_bullet(doc, "Open a Command Prompt and install the required packages:")
add_code_block(doc, ["pip install python-docx pymysql openpyxl"])
add_bullet(doc,
    "Verify that XAMPP (or equivalent) is running and the MySQL database is accessible.",
    bold_parts=["XAMPP"])
add_bullet(doc, "Confirm the FTP server and API endpoint are reachable from the workstation.")

add_step_heading(doc, "Step 2: Configure config.ini")
add_bullet(doc,
    "On first launch the application automatically creates a config.ini template in the "
    "program folder. Open it with any text editor and fill in the actual values.")
add_code_block(doc, [
    "[DB]",
    "host     = <MySQL server IP or hostname>",
    "port     = 3306",
    "database = <database name>",
    "username = <MySQL username>",
    "password = <MySQL password>",
    "",
    "[FTP]",
    "host       = <FTP server IP or hostname>",
    "port       = 21",
    "username   = <FTP username>",
    "password   = <FTP password>",
    "remote_dir = /attach",
    "",
    "[API]",
    "base_url = https://<your-server>/eclabapi",
])
add_bullet(doc, "Save the file. Configuration is loaded at every startup — no restart is needed "
                "if you update credentials.")
add_warning(doc,
    "Keep config.ini confidential. It contains database and FTP passwords. "
    "Do not share or commit this file.")

# ═══════════════════════════════════════════════════════════════════════════
# PART 2
# ═══════════════════════════════════════════════════════════════════════════
add_part_heading(doc, "PART 2 — Launching the System")

add_step_heading(doc, "Step 3: Open the Application")
add_bullet(doc,
    "Double-click epass_upload_gui.py (or the compiled .exe if provided) to launch the system.",
    bold_parts=["epass_upload_gui.py", ".exe"])
add_bullet(doc,
    "The main window titled “ePass TP Upload” opens, showing:",
    )
add_bullet(doc, "Two folder-selection rows: ePass Folder and GC Folder.")
add_bullet(doc, "A Process & Upload button (disabled until at least one folder is chosen).")
add_bullet(doc, "A dark log panel for real-time output.")
add_bullet(doc, "A Schedule... button for configuring automated runs.")

# ═══════════════════════════════════════════════════════════════════════════
# PART 3
# ═══════════════════════════════════════════════════════════════════════════
add_part_heading(doc, "PART 3 — Preparing the PDF Documents")

add_step_heading(doc, "Step 4: Prepare ePass PDF Files")
add_bullet(doc,
    "Collect all ePass PDF files for the current batch into one folder "
    "(e.g. C:\\Documents\\ePass_Today).")
add_bullet(doc,
    "Files can have any filename; the system renames them automatically using the "
    "worker’s passport number and name.")
add_bullet(doc,
    "Both supported ePass layouts (Format A and Format B) are detected automatically — "
    "no manual sorting is needed.")

add_step_heading(doc, "Step 5: Prepare Green Card (GC) PDF Files")
add_bullet(doc,
    "Collect all Green Card PDF files into a separate folder "
    "(e.g. C:\\Documents\\GC_Today).")
add_bullet(doc,
    "The system reads Malay-labelled fields (No. K.P., Nama Personel, Warganegara) "
    "to extract worker information.")
add_bullet(doc,
    "You may run the system with only ePass files, only GC files, or both simultaneously.")

# ═══════════════════════════════════════════════════════════════════════════
# PART 4
# ═══════════════════════════════════════════════════════════════════════════
add_part_heading(doc, "PART 4 — Running a Manual Upload")

add_step_heading(doc, "Step 6: Select PDF Folders")
add_bullet(doc,
    "Click Browse... next to ePass Folder: and navigate to the folder containing ePass PDFs.",
    bold_parts=["Browse...", "ePass Folder:"])
add_bullet(doc,
    "Click Browse... next to GC Folder: and navigate to the folder containing GC PDFs.",
    bold_parts=["Browse...", "GC Folder:"])
add_bullet(doc,
    "Selecting at least one folder activates the Process & Upload button.",
    bold_parts=["Process & Upload"])

add_step_heading(doc, "Step 7: Start Processing")
add_bullet(doc,
    "Click Process & Upload. The button is replaced by a progress bar.",
    bold_parts=["Process & Upload"])
add_bullet(doc,
    "The system performs the following automatically:")
add_bullet(doc,
    "Scans both folders and extracts worker details (passport number, name, nationality) from each PDF.")
add_bullet(doc,
    "Queries two listing APIs — api_epass.php (ePass status) and api_tp.php (TP/GC status) — "
    "and merges the results to identify which documents each worker still needs. "
    "A document is treated as pending if its API value is null, empty, or blank.")
add_bullet(doc,
    "Falls back to the workers database if both listing APIs are unavailable.")
add_bullet(doc,
    "Renames each PDF to the standard format: PASSPORT - NAME_ePass.pdf or PASSPORT - NAME_GC.pdf.",
    bold_parts=["PASSPORT - NAME_ePass.pdf", "PASSPORT - NAME_GC.pdf"])
add_bullet(doc,
    "Connects to the FTP server, creates a per-contractor directory under /attach/, and uploads the files.",
    bold_parts=["/attach/"])
add_bullet(doc,
    "Notifies the API endpoints that the upload is complete for each worker.")
add_warning(doc,
    "Do not close the application while processing is in progress. Doing so may leave files partially uploaded.")

# ═══════════════════════════════════════════════════════════════════════════
# PART 5
# ═══════════════════════════════════════════════════════════════════════════
add_part_heading(doc, "PART 5 — Reviewing Upload Results")

add_step_heading(doc, "Step 8: Read the On-Screen Log")
add_bullet(doc,
    "Once processing is complete, review the log panel for the session summary:")
add_code_block(doc, [
    "Completed in 12.5s",
    "  Workers:         5",
    "  Success:         4",
    "  Skipped:         1",
    "  API failures:    0",
    "  Upload failures: 0",
    "  Read errors:     0",
])
add_bullet(doc,
    "Common status values for individual workers:")
add_bullet(doc, "Success — ePass and/or GC uploaded and API notified.")
add_bullet(doc, "Skipped — Already complete — documents already recorded in the system.")
add_bullet(doc, "Skipped — Not in DB — passport number not found in the database.")
add_bullet(doc, "Uploaded / API Failed — files uploaded but API notification failed.")
add_bullet(doc, "Failed — FTP upload could not be completed.")

add_step_heading(doc, "Step 9: Open the Excel Report")
add_bullet(doc,
    "Click Open Excel Report to open the automatically generated spreadsheet.",
    bold_parts=["Open Excel Report"])
add_bullet(doc,
    "The report is saved to the Report\\ folder with a timestamped filename:",
    bold_parts=["Report\\"])
add_code_block(doc, ["Report\\EPassReport_YYYYMMDD_HHMMSS.xlsx"])
add_bullet(doc,
    "Each row represents one worker. Row colours indicate status:")
add_bullet(doc, "Green — upload successful.")
add_bullet(doc, "Red — upload or API failure.")
add_bullet(doc, "Light grey — skipped.")
add_bullet(doc,
    "A summary totals row at the bottom shows aggregate success, skipped, and failure counts.")

# ═══════════════════════════════════════════════════════════════════════════
# PART 6
# ═══════════════════════════════════════════════════════════════════════════
add_part_heading(doc, "PART 6 — Scheduling Automated Daily Runs")

add_step_heading(doc, "Step 10: Open the Scheduler Dialog")
add_bullet(doc,
    "Click the Schedule... button in the main window.",
    bold_parts=["Schedule..."])
add_bullet(doc,
    "The Schedule Daily Run dialog opens with the following fields:")
add_bullet(doc, "Run time (HH:MM) — the time the task will execute each day (24-hour format).")
add_bullet(doc, "ePass folder — full path to the folder containing ePass PDFs.")
add_bullet(doc, "GC folder — full path to the folder containing GC PDFs.")

add_step_heading(doc, "Step 11: Configure and Register the Task")
add_bullet(doc,
    "Enter the desired run time (e.g. 08:00 for 8 AM) and browse to or type the folder paths.")
add_bullet(doc,
    "Click Register Task. The system validates the time format and paths, then creates a "
    "Windows Task Scheduler entry named ePass TP Upload.",
    bold_parts=["Register Task", "ePass TP Upload"])
add_bullet(doc,
    "During a scheduled run the application launches in headless mode (no GUI window) "
    "and logs all output to Log\\upload.log.",
    bold_parts=["Log\\upload.log"])
add_bullet(doc,
    "If the source folders are empty, the run is still logged to Log\\upload.log "
    "but no Excel report is generated.",
    bold_parts=["Log\\upload.log"])
add_bullet(doc,
    "Use the buttons in the dialog to manage the task:")
add_bullet(doc, "Enable / Disable — temporarily pause scheduled runs without deleting the task.")
add_bullet(doc, "Remove Task — permanently delete the scheduled task.")
add_bullet(doc,
    "The Run History table in the dialog shows past scheduled runs with date, worker count, "
    "and success/failure totals.")
add_warning(doc,
    "Ensure the workstation remains powered on and the user account is logged in at the "
    "scheduled time, or configure the task to run whether or not the user is logged in "
    "via Windows Task Scheduler Advanced Settings.")

# ── PAGE BREAK ───────────────────────────────────────────────────────────────
add_page_break(doc)

# ── FOLDER & FILE REFERENCE TABLE ────────────────────────────────────────────
ref_heading = doc.add_paragraph()
set_spacing(ref_heading, before=280, after=40)
rh_run = ref_heading.add_run("Folder & File Structure Reference")
rh_run.bold = True
rh_run.font.color.rgb = NAVY

tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Table Grid"
tbl.rows[0].cells[0].width = Twips(3969)
tbl.rows[0].cells[1].width = Twips(5102)

# Header row
hdr_cells = tbl.rows[0].cells
for cell in hdr_cells:
    set_cell_bg(cell, NAVY_HEX)

hdr_cells[0].paragraphs[0].clear()
hdr_cells[1].paragraphs[0].clear()

def hdr_run(cell, text):
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(10)

hdr_run(hdr_cells[0], "Folder / File")
hdr_run(hdr_cells[1], "Purpose")

# Data rows
rows_data = [
    ("config.ini",                          "Database, FTP, and API credentials (edit before first use)"),
    ("epass_upload_gui.py",                 "Main application — run this to open the GUI"),
    ("Log\\upload.log",                     "Detailed log of all upload activity and API responses"),
    ("Report\\EPassReport_*.xlsx",          "Excel reports generated after each manual or scheduled run"),
    ("/attach/<CONTRACTOR_ID>/",            "FTP remote directory structure where files are stored"),
    ("<ePass Folder>\\*_ePass.pdf",         "Renamed ePass files (PASSPORT - NAME_ePass.pdf)"),
    ("<GC Folder>\\*_GC.pdf",              "Renamed Green Card files (PASSPORT - NAME_GC.pdf)"),
]

for i, (folder, purpose) in enumerate(rows_data):
    row_cells = tbl.add_row().cells
    bg = WHITE_HEX if i % 2 == 0 else LGRAY_HEX
    for cell in row_cells:
        set_cell_bg(cell, bg)

    # Folder cell — Courier New
    row_cells[0].paragraphs[0].clear()
    run0 = row_cells[0].paragraphs[0].add_run(folder)
    run0.font.name = "Courier New"
    run0.font.size = Pt(10)

    # Purpose cell — regular
    row_cells[1].paragraphs[0].clear()
    run1 = row_cells[1].paragraphs[0].add_run(purpose)
    run1.font.size = Pt(10)

# ── FOOTER LINE ───────────────────────────────────────────────────────────────
doc.add_paragraph()
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(footer_p, before=0, after=0)
foot_run = footer_p.add_run("For technical support, contact your IT department.")
foot_run.italic = True
foot_run.font.color.rgb = GRAY
foot_run.font.size = Pt(9)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out_path = r"C:\xampp\htdocs\epass_tp_upload\ePassTPUpload_Tutorial.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
